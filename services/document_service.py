"""Turning an uploaded file into chunks that can be cited.

The whole point of this module is the **page number**. Extracting text from a PDF is easy;
extracting it while remembering which page each character came from is the part that makes a
citation checkable by a human. Everything here is arranged around not losing that.

Chunking strategy: fixed-size windows with overlap, split on paragraph and sentence boundaries
where one is available nearby. Fixed windows are predictable and easy to explain; the boundary
search stops a chunk ending mid-sentence, which is what makes a retrieved snippet readable when
it is shown as a citation. Phase 8's chunk-size experiment measures what the size costs.
"""
from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass
from pathlib import Path

from core.config import settings
from core.logging import get_logger

log = get_logger("documents")

SUPPORTED_TYPES = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
}


class UnsupportedDocument(Exception):
    """The file extension is not one this platform can read."""


class EmptyDocument(Exception):
    """The file parsed successfully but contained no extractable text."""


@dataclass
class Page:
    """One page of extracted text. Non-paginated formats produce a single page."""

    number: int
    text: str


@dataclass
class Chunk:
    ordinal: int
    text: str
    page: int | None
    char_start: int
    char_end: int


# ------------------------------------------------------------------- extraction
def _clean(text: str) -> str:
    """Normalise whitespace without destroying paragraph structure.

    PDFs in particular arrive full of soft hyphens, non-breaking spaces and ligatures. Left in,
    they break both keyword search (``ﬁle`` does not match ``file``) and the readability of a
    quoted snippet.
    """
    text = unicodedata.normalize("NFKC", text)
    text = text.replace("­", "")                 # soft hyphen
    text = re.sub(r"[ \t\f\v]+", " ", text)           # runs of spaces, but not newlines
    text = re.sub(r"\n{3,}", "\n\n", text)            # at most one blank line
    return text.strip()


def extract_pages(path: Path, suffix: str) -> list[Page]:
    """Text per page. Formats without pages return one page numbered 1."""
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = _clean(page.extract_text() or "")
            if text:
                pages.append(Page(number=index, text=text))
        return pages

    if suffix == ".docx":
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        # Tables carry a lot of the substance in reports, and python-docx keeps them out of
        # `paragraphs`. Dropping them silently loses content the user can see in Word.
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = _clean("\n\n".join(parts))
        return [Page(number=1, text=text)] if text else []

    if suffix in {".txt", ".md", ".markdown"}:
        raw = path.read_text(encoding="utf-8", errors="replace")
        text = _clean(raw)
        return [Page(number=1, text=text)] if text else []

    raise UnsupportedDocument(suffix)


# --------------------------------------------------------------------- chunking
# How far back to look for a clean break rather than cutting mid-sentence. A quarter of the
# chunk: far enough to usually find one, near enough that chunks stay roughly the target size.
def _boundary_window(chunk_size: int) -> int:
    return max(80, chunk_size // 4)


def _split_point(text: str, start: int, end: int, chunk_size: int) -> int:
    """The best place to end a chunk that begins at ``start``.

    Preference order: a paragraph break, then a sentence end, then a space, then the hard limit.
    """
    if end >= len(text):
        return len(text)

    window_start = max(start + 1, end - _boundary_window(chunk_size))
    window = text[window_start:end]

    paragraph = window.rfind("\n\n")
    if paragraph != -1:
        return window_start + paragraph + 2

    sentence = max(window.rfind(". "), window.rfind("? "), window.rfind("! "), window.rfind(".\n"))
    if sentence != -1:
        return window_start + sentence + 2

    space = window.rfind(" ")
    if space != -1:
        return window_start + space + 1

    return end


def chunk_pages(
    pages: list[Page],
    chunk_size: int | None = None,
    overlap: int | None = None,
) -> list[Chunk]:
    """Split pages into overlapping chunks, each remembering the page it came from.

    Chunking happens **within** a page rather than across the whole document, so every chunk has
    exactly one page number. A chunk spanning a page break would have to cite two pages or lie
    about one, and a citation that is only mostly right is worse than one that is narrower.
    """
    chunk_size = chunk_size or settings.chunk_size
    overlap = overlap if overlap is not None else settings.chunk_overlap
    # Overlap at or above the chunk size would never advance the cursor.
    overlap = min(overlap, chunk_size // 2)

    chunks: list[Chunk] = []
    ordinal = 0

    for page in pages:
        text = page.text
        position = 0
        while position < len(text):
            hard_end = min(position + chunk_size, len(text))
            end = _split_point(text, position, hard_end, chunk_size)
            body = text[position:end].strip()

            if body:
                chunks.append(
                    Chunk(
                        ordinal=ordinal,
                        text=body,
                        page=page.number,
                        char_start=position,
                        char_end=end,
                    )
                )
                ordinal += 1

            if end >= len(text):
                break
            position = max(end - overlap, position + 1)

    return chunks


# ------------------------------------------------------------------- validation
def validate_upload(filename: str, size_bytes: int) -> tuple[str, str]:
    """Check the extension and size. Returns ``(suffix, mime_type)``."""
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_TYPES:
        raise UnsupportedDocument(
            f"{suffix or 'that file type'} is not supported. "
            f"Upload a PDF, DOCX, TXT or Markdown file."
        )
    limit = settings.max_upload_mb * 1024 * 1024
    if size_bytes > limit:
        raise UnsupportedDocument(
            f"That file is {size_bytes / 1_048_576:.1f} MB. The limit is {settings.max_upload_mb} MB."
        )
    return suffix, SUPPORTED_TYPES[suffix]


def parse(path: Path, suffix: str) -> tuple[list[Page], list[Chunk]]:
    """Extract and chunk in one step. Raises :class:`EmptyDocument` if nothing readable."""
    pages = extract_pages(path, suffix)
    if not pages:
        raise EmptyDocument(
            "No text could be read from that file. Scanned PDFs need OCR, which this platform "
            "does not do yet."
        )
    chunks = chunk_pages(pages)
    if not chunks:
        raise EmptyDocument("The file parsed but produced no usable text.")
    log.info("Parsed %s: %d pages, %d chunks", path.name, len(pages), len(chunks))
    return pages, chunks
